#!/usr/bin/env python3
"""Verify canonical report facts and their HTML/DOCX/PDF representations."""
from __future__ import annotations

import argparse
import json
import re
import subprocess
import sys
from pathlib import Path

from docx import Document


def _read_docx(path: Path) -> str:
    document = Document(path)
    return "\n".join([*(p.text for p in document.paragraphs), *(cell.text for table in document.tables for row in table.rows for cell in row.cells)])


def _read_pdf(path: Path) -> str:
    completed = subprocess.run(["pdftotext", str(path), "-"], text=True, capture_output=True, check=False)
    if completed.returncode:
        raise RuntimeError(f"pdftotext failed: {completed.stderr.strip()}")
    return completed.stdout


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--case", type=Path, required=True)
    parser.add_argument("--report-json", type=Path, required=True)
    parser.add_argument("--web-html", type=Path, required=True)
    parser.add_argument("--docx", type=Path, required=True)
    parser.add_argument("--pdf", type=Path, required=True)
    parser.add_argument("--output", type=Path)
    args = parser.parse_args()
    model = json.loads(args.report_json.read_text(encoding="utf-8"))
    texts = {"web": args.web_html.read_text(encoding="utf-8"), "docx": _read_docx(args.docx), "pdf": _read_pdf(args.pdf)}
    item_names = [row["original_name"] for row in model["service_catalog"]]
    required = [str(model["metadata"]["procurement_number"]), model["procurement_passport"]["okpd2"], "Требуется дополнительная проверка", "Проект контракта отсутствует", *item_names]
    forbidden = ("смэв", "интеграц", "обучени", "преподавател", "аудитори", "undefined", "lorem ipsum")
    normalized = {name: re.sub(r"\s+", " ", text).strip() for name, text in texts.items()}
    format_failures = {name: [term for term in required if term not in text] + [term for term in forbidden if term in text.lower()] for name, text in normalized.items()}
    gates = {
        "canonical_service_rows": len(model["service_catalog"]) == 43,
        "canonical_decision": model["bid_decision"]["status"] == "needs_review",
        "canonical_no_line_totals": all(row["line_total"] is None for row in model["service_catalog"]),
        "canonical_unknown_quantity": all(row["quantity"] is None for row in model["service_catalog"]),
        "format_parity": not any(format_failures.values()),
        "evidence_coverage": len(model["evidence_map"]) == len(model["service_catalog"]),
    }
    result = {"status": "passed" if all(gates.values()) else "failed", "gates": gates, "format_failures": format_failures}
    payload = json.dumps(result, ensure_ascii=False, indent=2) + "\n"
    if args.output:
        args.output.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(payload, encoding="utf-8")
    print(payload, end="")
    return 0 if result["status"] == "passed" else 1


if __name__ == "__main__":
    sys.exit(main())
