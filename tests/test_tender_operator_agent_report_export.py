from __future__ import annotations

from pathlib import Path
from concurrent.futures import ThreadPoolExecutor
from unittest.mock import patch

import pytest
from fastapi import HTTPException

from src.modules.tender_operator_agent_demo.report_export_service import (
    ExportedDemoReport,
    _pdf_artifact_paths,
    export_demo_agent_report_docx,
    export_demo_agent_report_pdf,
)
from src.modules.tender_operator_agent_demo.schemas import TenderOperatorDemoReportResponse


SAMPLE_METADATA = {
    "run_id": "toa-run-test-00000000-abc123",
    "created_at": "2026-07-06T12:00:00+00:00",
    "mode": "procurement_search_intake",
    "tender_title": "Тестовая закупка",
    "tender_category": "44-ФЗ",
    "customer_name": "ООО Тестовый Заказчик",
    "status": "needs_review",
    "analysis_mode": "fallback_deterministic_adapter",
    "procurement_source": "zakupki_gov_ru_getdocs_ip",
    "procurement_id": "0123456789012345",
    "procurement_url": "https://zakupki.gov.ru/example",
    "publication_date": "06.07.2026",
    "deadline": "20.07.2026",
    "notice_source_label": "электронное извещение ЕИС",
    "procurement": {
        "initial_price": 1234567.89,
        "currency": "RUB",
        "structured_source_label": "электронное извещение ЕИС",
    },
}

SAMPLE_REPORT_MARKDOWN = """# Отчёт по загруженному прогону тендерного агента

- Run ID: toa-run-test-00000000-abc123
- Закупка: Тестовая закупка
- Категория: 44-ФЗ
- Заказчик: ООО Тестовый Заказчик
- Статус: needs_review
- Режим анализа: fallback_deterministic_adapter

## Источник закупки
- Источник: zakupki_gov_ru_getdocs_ip
- Номер извещения: 0123456789012345
- Заказчик: ООО Тестовый Заказчик
- Закон: 44-ФЗ
- НМЦК: 1234567.89 RUB
- Срок подачи: 20.07.2026
- Источник сведений: электронное извещение ЕИС
- Статус скачивания: downloaded
- Ручная загрузка требовалась: нет
- Скачано/добавлено файлов: 1

### Документация
- documentation-archive.zip

## Краткий вывод
- Тестовый вывод.
- Анализ выполнен.

## Предварительный анализ закупки
### Ключевые требования и ограничения
- Требуется ручная валидация.

## Экономика
- Минимальная закупочная стоимость: 1000000.00
- Целевая маржа: 15%

## Ручные проверки
- Тестовая проверка 1.
- Тестовая проверка 2.
"""


@pytest.fixture(autouse=True)
def _clean_exports(tmp_path):
    with patch("src.modules.tender_operator_agent_demo.report_export_service._safe_output_dir", return_value=tmp_path):
        yield


@pytest.fixture
def mock_run_data():
    with (
        patch("src.modules.tender_operator_agent_demo.report_export_service._load_metadata", return_value=SAMPLE_METADATA),
        patch(
            "src.modules.tender_operator_agent_demo.report_export_service.get_uploaded_demo_report",
            return_value=TenderOperatorDemoReportResponse(
                run_id="toa-run-test-00000000-abc123",
                report_title="Отчёт по загруженному прогону тендерного агента",
                generated_at="2026-07-06T12:00:00",
                recommendation="manual_review_required",
                recommendation_label="нужна ручная проверка",
                executive_summary=["Тестовый вывод.", "Анализ выполнен."],
                manual_checks=["Тестовая проверка 1.", "Тестовая проверка 2."],
                sections=[],
                report_markdown=SAMPLE_REPORT_MARKDOWN,
            ),
        ),
        patch(
            "src.modules.tender_operator_agent_demo.report_export_service.get_uploaded_demo_report_html",
            return_value="<h1>Отчёт</h1><p>НМЦК: 1 234 567,89 RUB</p><p>Дата публикации: 06.07.2026</p><p>Срок подачи: 20.07.2026</p><p>Источник сведений: электронное извещение ЕИС</p>",
        ),
    ):
        yield


class TestDocxExport:
    def test_export_docx_creates_file(self, mock_run_data, tmp_path):
        result = export_demo_agent_report_docx("toa-run-test-00000000-abc123")
        assert result.format == "docx"
        assert result.file_name.endswith(".docx")
        path = Path(result.file_path)
        assert path.exists()
        assert path.suffix == ".docx"
        assert path.stat().st_size > 0

    def test_export_docx_content_type(self, mock_run_data, tmp_path):
        result = export_demo_agent_report_docx("toa-run-test-00000000-abc123")
        assert "wordprocessingml.document" in result.content_type

    def test_export_docx_filename_pattern(self, mock_run_data, tmp_path):
        result = export_demo_agent_report_docx("toa-run-test-00000000-abc123")
        assert result.file_name.startswith("demo_agent_report_")
        assert "0123456789012345" in result.file_name

    def test_export_docx_contains_key_fields(self, mock_run_data, tmp_path):
        result = export_demo_agent_report_docx("toa-run-test-00000000-abc123")
        from docx import Document
        doc = Document(result.file_path)
        text = " ".join(p.text for p in doc.paragraphs)
        assert "0123456789012345" in text
        assert "электронное извещение ЕИС" in text
        assert "1234567.89" in text or "1 234 567" in text
        assert "06.07.2026" in text
        assert "20.07.2026" in text

    def test_export_docx_registry_number_and_run_id(self, mock_run_data, tmp_path):
        result = export_demo_agent_report_docx("toa-run-test-00000000-abc123")
        assert result.registry_number == "0123456789012345"
        assert result.run_id == "toa-run-test-00000000-abc123"


class TestPdfExport:
    def test_export_pdf_creates_file(self, mock_run_data, tmp_path):
        result = export_demo_agent_report_pdf("toa-run-test-00000000-abc123")
        assert result.format == "pdf"
        assert result.file_name.endswith(".pdf")
        path = Path(result.file_path)
        assert path.exists()
        assert path.suffix == ".pdf"
        assert path.stat().st_size > 0

    def test_export_pdf_content_type(self, mock_run_data, tmp_path):
        result = export_demo_agent_report_pdf("toa-run-test-00000000-abc123")
        assert result.content_type == "application/pdf"

    def test_export_pdf_magic_bytes(self, mock_run_data, tmp_path):
        result = export_demo_agent_report_pdf("toa-run-test-00000000-abc123")
        with open(result.file_path, "rb") as f:
            header = f.read(5)
        assert header == b"%PDF-"

    def test_export_pdf_filename_pattern(self, mock_run_data, tmp_path):
        result = export_demo_agent_report_pdf("toa-run-test-00000000-abc123")
        assert result.file_name.startswith("demo_agent_report_")
        assert "0123456789012345" in result.file_name

    def test_export_pdf_has_valid_size_and_metadata(self, mock_run_data, tmp_path):
        result = export_demo_agent_report_pdf("toa-run-test-00000000-abc123")
        assert result.file_name.startswith("demo_agent_report_")
        assert "0123456789012345" in result.file_name
        assert result.file_name.endswith(".pdf")
        assert Path(result.file_path).stat().st_size > 1000

    def test_concurrent_first_export_publishes_one_valid_artifact(self, mock_run_data):
        with ThreadPoolExecutor(max_workers=4) as pool:
            exports = list(pool.map(lambda _: export_demo_agent_report_pdf("toa-run-test-00000000-abc123"), range(4)))
        payloads = {Path(item.file_path).read_bytes() for item in exports}
        assert len(payloads) == 1
        assert next(iter(payloads)).startswith(b"%PDF-")

    def test_corrupt_nonempty_artifact_is_rejected(self, mock_run_data, tmp_path):
        result = export_demo_agent_report_pdf("toa-run-test-00000000-abc123")
        path = Path(result.file_path)
        path.write_bytes(b"not-a-pdf")
        with pytest.raises(RuntimeError, match="invalid"):
            export_demo_agent_report_pdf("toa-run-test-00000000-abc123")

    def test_same_prefix_runs_have_distinct_immutable_pdf_artifacts(self, tmp_path):
        first = "toa-run-20260721144956-8b2eee"
        second = "toa-run-20260722110000-abcd12"
        first_metadata = {**SAMPLE_METADATA, "run_id": first}
        second_metadata = {**SAMPLE_METADATA, "run_id": second}
        reports = {
            first: TenderOperatorDemoReportResponse(run_id=first, report_title="one", generated_at="2026-07-06T12:00:00", recommendation="manual_review_required", recommendation_label="ok", executive_summary=[], manual_checks=[], sections=[], report_markdown="first report"),
            second: TenderOperatorDemoReportResponse(run_id=second, report_title="two", generated_at="2026-07-06T12:00:00", recommendation="manual_review_required", recommendation_label="ok", executive_summary=[], manual_checks=[], sections=[], report_markdown="second report"),
        }
        with (
            patch("src.modules.tender_operator_agent_demo.report_export_service._load_metadata", side_effect=lambda run_id: first_metadata if run_id == first else second_metadata),
            patch("src.modules.tender_operator_agent_demo.report_export_service.get_uploaded_demo_report", side_effect=lambda run_id: reports[run_id]),
            patch("src.modules.tender_operator_agent_demo.report_export_service.get_uploaded_demo_report_html", return_value="<p>report</p>"),
        ):
            first_export = export_demo_agent_report_pdf(first)
            second_export = export_demo_agent_report_pdf(second)
            assert first_export.file_path != second_export.file_path
            assert Path(first_export.file_path).read_bytes() != Path(second_export.file_path).read_bytes()
            assert export_demo_agent_report_pdf(first).file_path == first_export.file_path
            assert export_demo_agent_report_pdf(second).file_path == second_export.file_path
            assert __import__("json").loads(Path(first_export.file_path).with_suffix(".manifest.json").read_text())["run_id"] == first
            assert __import__("json").loads(Path(second_export.file_path).with_suffix(".manifest.json").read_text())["run_id"] == second


class TestErrors:
    def test_unknown_run_id_raises_not_found(self):
        with patch("src.modules.tender_operator_agent_demo.report_export_service._load_metadata") as mock_load:
            mock_load.side_effect = HTTPException(status_code=404, detail="Run not found")
            with pytest.raises(HTTPException, match="Run not found"):
                export_demo_agent_report_docx("nonexistent-run")

    def test_report_not_available(self):
        with (
            patch("src.modules.tender_operator_agent_demo.report_export_service._load_metadata", return_value=SAMPLE_METADATA),
            patch("src.modules.tender_operator_agent_demo.report_export_service.get_uploaded_demo_report") as mock_report,
        ):
            mock_report.side_effect = HTTPException(status_code=404, detail="Report is not available yet")
            with pytest.raises(HTTPException, match="Report is not available yet"):
                export_demo_agent_report_docx("toa-run-test-00000000-abc123")

    def test_invalid_run_id_rejected_before_file_access(self):
        with pytest.raises(ValueError, match="Invalid run ID"):
            export_demo_agent_report_docx("../secret")


class TestHtmlFallback:
    def test_export_docx_uses_html_when_markdown_missing(self, tmp_path):
        with (
            patch("src.modules.tender_operator_agent_demo.report_export_service._safe_output_dir", return_value=tmp_path),
            patch("src.modules.tender_operator_agent_demo.report_export_service._load_metadata", return_value=SAMPLE_METADATA),
            patch(
                "src.modules.tender_operator_agent_demo.report_export_service.get_uploaded_demo_report",
                return_value=TenderOperatorDemoReportResponse(
                    run_id="toa-run-test-00000000-abc123",
                    report_title="Отчёт",
                    generated_at="2026-07-06T12:00:00",
                    recommendation="manual_review_required",
                    recommendation_label="нужна ручная проверка",
                    executive_summary=[],
                    manual_checks=[],
                    sections=[],
                    report_markdown="",
                ),
            ),
            patch(
                "src.modules.tender_operator_agent_demo.report_export_service.get_uploaded_demo_report_html",
                return_value="<h1>Отчёт</h1><p>НМЦК: 1 234 567,89 RUB</p><p>Дата публикации: 06.07.2026</p><p>Срок подачи: 20.07.2026</p><ul><li>Источник сведений: электронное извещение ЕИС</li></ul>",
            ),
        ):
            result = export_demo_agent_report_docx("toa-run-test-00000000-abc123")
            from docx import Document

            doc = Document(result.file_path)
            text = " ".join(p.text for p in doc.paragraphs)
            assert "1 234 567,89" in text or "1 234 567" in text
            assert "06.07.2026" in text
            assert "20.07.2026" in text
            assert "электронное извещение ЕИС" in text


class TestApi:
    def test_demo_run_export_docx_api_returns_file(self, client, tmp_path):
        docx_path = tmp_path / "demo.docx"
        docx_path.write_bytes(b"PK\x03\x04demo-docx")
        exported = ExportedDemoReport(
            run_id="toa-run-test-00000000-abc123",
            registry_number="0123456789012345",
            format="docx",
            file_name="demo.docx",
            file_path=str(docx_path),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            size_bytes=123,
            created_at="2026-07-06T12:00:00+00:00",
            source_report_path=None,
        )
        with patch("src.modules.tender_operator_agent_demo.router.export_demo_agent_report_docx", return_value=exported):
            response = client.get("/api/demo/tender-agent/runs/toa-run-test-00000000-abc123/export/docx")
        assert response.status_code == 200
        assert "wordprocessingml.document" in response.headers["content-type"]
        assert "filename=\"demo.docx\"" in response.headers["content-disposition"]
        assert response.content.startswith(b"PK")

    def test_demo_run_export_pdf_api_returns_file(self, client, tmp_path):
        pdf_path = tmp_path / "demo.pdf"
        pdf_path.write_bytes(b"%PDF-1.4\n%demo\n")
        exported = ExportedDemoReport(
            run_id="toa-run-test-00000000-abc123",
            registry_number="0123456789012345",
            format="pdf",
            file_name="demo.pdf",
            file_path=str(pdf_path),
            content_type="application/pdf",
            size_bytes=pdf_path.stat().st_size,
            created_at="2026-07-06T12:00:00+00:00",
            source_report_path=None,
        )
        with patch("src.modules.tender_operator_agent_demo.router.export_demo_agent_report_pdf", return_value=exported):
            response = client.get("/api/demo/tender-agent/runs/toa-run-test-00000000-abc123/export/pdf")
        assert response.status_code == 200
        assert response.headers["content-type"].startswith("application/pdf")
        assert response.content.startswith(b"%PDF-")

    def test_demo_run_export_unknown_run_returns_404(self, client):
        with patch(
            "src.modules.tender_operator_agent_demo.router.export_demo_agent_report_docx",
            side_effect=FileNotFoundError("Run was not found"),
        ):
            response = client.get("/api/demo/tender-agent/runs/nonexistent/export/docx")
        assert response.status_code == 404

    def test_demo_run_export_report_missing_returns_404(self, client):
        with patch(
            "src.modules.tender_operator_agent_demo.router.export_demo_agent_report_pdf",
            side_effect=HTTPException(status_code=404, detail="Report is not available yet"),
        ):
            response = client.get("/api/demo/tender-agent/runs/toa-run-test-00000000-abc123/export/pdf")
        assert response.status_code == 404
