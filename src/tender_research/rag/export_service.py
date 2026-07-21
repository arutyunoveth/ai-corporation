from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime, timezone
from html import escape as html_escape
from pathlib import Path
import re
from typing import Literal
from xml.sax.saxutils import escape as xml_escape

from docx import Document
from docx.shared import Pt
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer
from sqlalchemy.orm import Session

from src.tender_research.config import load_config
from src.tender_research.rag.history_service import AnalysisRunRecord, get_analysis_run, get_analysis_run_report

DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
PDF_CONTENT_TYPE = "application/pdf"
_EXPORT_SUBDIR = ("rag", "exports")
_FONT_NAME = "TenderExportUnicode"
_FONT_CANDIDATES = (
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    "/usr/share/fonts/dejavu/DejaVuSans.ttf",
    "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    "/System/Library/Fonts/Supplemental/Arial Unicode MS.ttf",
    "/System/Library/Fonts/Supplemental/Arial.ttf",
    "/System/Library/Fonts/Supplemental/Verdana.ttf",
    "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
)


@dataclass(frozen=True)
class ExportedReport:
    analysis_run_id: str
    registry_number: str
    format: Literal["docx", "pdf"]
    file_name: str
    file_path: str
    content_type: str
    size_bytes: int
    created_at: str
    source_report_path: str | None


@dataclass(frozen=True)
class _MarkdownBlock:
    kind: Literal["heading", "paragraph", "bullet_list", "numbered_list"]
    text: str = ""
    level: int = 0
    items: tuple[str, ...] = ()


def _safe_segment(value: str, fallback: str) -> str:
    cleaned = re.sub(r"[^0-9A-Za-z_-]+", "_", (value or "").strip()).strip("_")
    return cleaned or fallback


def _short_run_id(run_id: str) -> str:
    return _safe_segment(run_id, "run")[:8]


def _analysis_title(record: AnalysisRunRecord) -> str:
    return "Анализ закупки"


def _build_export_file_name(registry_number: str, analysis_run_id: str, format_name: Literal["docx", "pdf"]) -> str:
    safe_registry = _safe_segment(registry_number, "unknown_registry")
    return f"tender_analysis_{safe_registry}_{_short_run_id(analysis_run_id)}.{format_name}"


def _safe_output_dir(output_dir: Path | None, *, data_dir: str) -> Path:
    root = (output_dir or Path(data_dir).joinpath(*_EXPORT_SUBDIR)).resolve()
    root.mkdir(parents=True, exist_ok=True)
    return root


def _safe_output_path(root: Path, file_name: str) -> Path:
    candidate = (root / file_name).resolve()
    try:
        candidate.relative_to(root)
    except ValueError as exc:
        raise ValueError("Export output path escapes controlled directory") from exc
    return candidate


def _load_saved_report(
    analysis_run_id: str,
    *,
    session: Session,
    data_dir: str,
) -> tuple[AnalysisRunRecord, str]:
    record, markdown, error = get_analysis_run_report(session, analysis_run_id, data_dir)
    if record is None:
        raise FileNotFoundError(error or "Analysis run not found")
    if markdown is None:
        raise FileNotFoundError(error or "Report file is missing or inaccessible")
    return record, markdown


def _metadata_lines(record: AnalysisRunRecord, markdown: str) -> list[str]:
    metadata = record.metadata or {}
    lines = [
        f"Реестровый номер: {record.registry_number}",
        f"Analysis run ID: {record.id}",
        f"Статус: {record.status}",
        f"Разделов: {record.sections_count}",
        f"Источников: {record.sources_count}",
        f"LLM: {'да' if record.used_llm else 'нет'}",
    ]
    if record.created_at:
        lines.append(f"Создано: {record.created_at.isoformat()}")
    if record.duration_seconds is not None:
        lines.append(f"Длительность: {record.duration_seconds:.2f} сек")
    if record.source:
        lines.append(f"Источник запуска: {record.source}")
    if record.llm_model:
        lines.append(f"LLM model: {record.llm_model}")
    if record.retrieval_provider or record.retrieval_model:
        retrieval_value = " / ".join(part for part in [record.retrieval_provider, record.retrieval_model] if part)
        lines.append(f"Retrieval: {retrieval_value}")
    if isinstance(metadata, dict):
        analysis_mode = metadata.get("analysis_mode")
        if analysis_mode:
            lines.append(f"Режим анализа: {analysis_mode}")
    warnings = record.warnings or []
    if warnings:
        lines.append("Warnings:")
        lines.extend([f"- {warning}" for warning in warnings])
    errors = record.errors or []
    if errors:
        lines.append("Errors:")
        lines.extend([f"- {item}" for item in errors])
    if markdown.strip():
        lines.append("")
    return lines


def _parse_markdown_blocks(markdown: str) -> list[_MarkdownBlock]:
    blocks: list[_MarkdownBlock] = []
    lines = markdown.splitlines()
    index = 0
    while index < len(lines):
        raw_line = lines[index].rstrip()
        line = raw_line.strip()
        if not line:
            index += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            blocks.append(_MarkdownBlock(kind="heading", level=min(len(heading_match.group(1)), 3), text=heading_match.group(2).strip()))
            index += 1
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)$", line)
        if bullet_match:
            items: list[str] = []
            while index < len(lines):
                current = lines[index].strip()
                match = re.match(r"^[-*]\s+(.*)$", current)
                if not match:
                    break
                items.append(match.group(1).strip())
                index += 1
            blocks.append(_MarkdownBlock(kind="bullet_list", items=tuple(items)))
            continue

        numbered_match = re.match(r"^\d+[.)]\s+(.*)$", line)
        if numbered_match:
            items = []
            while index < len(lines):
                current = lines[index].strip()
                match = re.match(r"^\d+[.)]\s+(.*)$", current)
                if not match:
                    break
                items.append(match.group(1).strip())
                index += 1
            blocks.append(_MarkdownBlock(kind="numbered_list", items=tuple(items)))
            continue

        paragraph_lines = [line]
        index += 1
        while index < len(lines):
            next_line = lines[index].strip()
            if not next_line:
                index += 1
                break
            if re.match(r"^(#{1,6})\s+", next_line) or re.match(r"^[-*]\s+", next_line) or re.match(r"^\d+[.)]\s+", next_line):
                break
            paragraph_lines.append(next_line)
            index += 1
        blocks.append(_MarkdownBlock(kind="paragraph", text=" ".join(paragraph_lines)))
    return blocks


def _iter_inline_parts(text: str) -> list[tuple[str, bool]]:
    parts: list[tuple[str, bool]] = []
    pattern = re.compile(r"\*\*(.+?)\*\*")
    cursor = 0
    for match in pattern.finditer(text):
        if match.start() > cursor:
            parts.append((text[cursor:match.start()], False))
        parts.append((match.group(1), True))
        cursor = match.end()
    if cursor < len(text):
        parts.append((text[cursor:], False))
    return parts or [(text, False)]


def _docx_add_formatted_paragraph(document: Document, text: str, *, style: str | None = None) -> None:
    paragraph = document.add_paragraph(style=style)
    for chunk, bold in _iter_inline_parts(text):
        run = paragraph.add_run(chunk)
        run.bold = bold


def _build_docx(
    record: AnalysisRunRecord,
    markdown: str,
    output_path: Path,
) -> None:
    document = Document()
    document.core_properties.title = f"{_analysis_title(record)} {record.registry_number}"
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10.5)

    title = document.add_paragraph()
    title.style = document.styles["Title"]
    title_run = title.add_run(_analysis_title(record))
    title_run.font.name = "Arial"

    for meta_line in _metadata_lines(record, markdown):
        if meta_line:
            _docx_add_formatted_paragraph(document, meta_line)
        else:
            document.add_paragraph()

    document.add_heading("Отчёт", level=1)

    for block in _parse_markdown_blocks(markdown):
        if block.kind == "heading":
            document.add_heading(block.text, level=block.level)
        elif block.kind == "paragraph":
            _docx_add_formatted_paragraph(document, block.text)
        elif block.kind == "bullet_list":
            for item in block.items:
                _docx_add_formatted_paragraph(document, item, style="List Bullet")
        elif block.kind == "numbered_list":
            for item in block.items:
                _docx_add_formatted_paragraph(document, item, style="List Number")

    document.save(output_path)


def _resolve_pdf_font_path() -> Path | None:
    for candidate in _FONT_CANDIDATES:
        path = Path(candidate)
        if path.exists():
            return path
    return None


def _ensure_pdf_font_registered() -> str:
    if _FONT_NAME in pdfmetrics.getRegisteredFontNames():
        return _FONT_NAME
    font_path = _resolve_pdf_font_path()
    if font_path is None:
        return "Helvetica"
    pdfmetrics.registerFont(TTFont(_FONT_NAME, str(font_path)))
    return _FONT_NAME


def _pdf_styles(font_name: str) -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "TenderExportTitle",
            parent=base["Title"],
            fontName=font_name,
            fontSize=20,
            leading=24,
            textColor=HexColor("#111111"),
            spaceAfter=10,
        ),
        "meta": ParagraphStyle(
            "TenderExportMeta",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=10,
            leading=14,
            alignment=TA_LEFT,
            spaceAfter=4,
        ),
        "h1": ParagraphStyle(
            "TenderExportH1",
            parent=base["Heading1"],
            fontName=font_name,
            fontSize=16,
            leading=20,
            spaceBefore=12,
            spaceAfter=6,
        ),
        "h2": ParagraphStyle(
            "TenderExportH2",
            parent=base["Heading2"],
            fontName=font_name,
            fontSize=13,
            leading=17,
            spaceBefore=10,
            spaceAfter=4,
        ),
        "h3": ParagraphStyle(
            "TenderExportH3",
            parent=base["Heading3"],
            fontName=font_name,
            fontSize=11,
            leading=15,
            spaceBefore=8,
            spaceAfter=4,
        ),
        "body": ParagraphStyle(
            "TenderExportBody",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=10.5,
            leading=14,
            spaceAfter=6,
        ),
    }


def _pdf_inline_markup(text: str) -> str:
    escaped = xml_escape(text)
    return re.sub(r"\*\*(.+?)\*\*", lambda match: f"<b>{xml_escape(match.group(1))}</b>", escaped)


def _build_pdf(record: AnalysisRunRecord, markdown: str, output_path: Path) -> None:
    font_name = _ensure_pdf_font_registered()
    styles = _pdf_styles(font_name)
    story = [
        Paragraph(html_escape(_analysis_title(record)), styles["title"]),
        Spacer(1, 2 * mm),
    ]

    for meta_line in _metadata_lines(record, markdown):
        if meta_line:
            story.append(Paragraph(_pdf_inline_markup(meta_line), styles["meta"]))
        else:
            story.append(Spacer(1, 2 * mm))

    story.append(Paragraph("Отчёт", styles["h1"]))

    for block in _parse_markdown_blocks(markdown):
        if block.kind == "heading":
            story.append(Paragraph(_pdf_inline_markup(block.text), styles[f"h{block.level}"]))
        elif block.kind == "paragraph":
            story.append(Paragraph(_pdf_inline_markup(block.text), styles["body"]))
        elif block.kind == "bullet_list":
            items = [ListItem(Paragraph(_pdf_inline_markup(item), styles["body"])) for item in block.items]
            story.append(ListFlowable(items, bulletType="bullet", start="circle"))
        elif block.kind == "numbered_list":
            items = [ListItem(Paragraph(_pdf_inline_markup(item), styles["body"])) for item in block.items]
            story.append(ListFlowable(items, bulletType="1"))

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=f"{_analysis_title(record)} {record.registry_number}",
    )
    doc.build(story)


def _build_exported_report(record: AnalysisRunRecord, output_path: Path, format_name: Literal["docx", "pdf"]) -> ExportedReport:
    return ExportedReport(
        analysis_run_id=record.id,
        registry_number=record.registry_number,
        format=format_name,
        file_name=output_path.name,
        file_path=str(output_path),
        content_type=DOCX_CONTENT_TYPE if format_name == "docx" else PDF_CONTENT_TYPE,
        size_bytes=output_path.stat().st_size,
        created_at=datetime.now(timezone.utc).isoformat(),
        source_report_path=record.report_path,
    )


def export_analysis_report_docx(
    analysis_run_id: str,
    *,
    output_dir: Path | None = None,
    data_dir: str | None = None,
    session: Session | None = None,
) -> ExportedReport:
    config_data_dir = data_dir or load_config().data_dir
    if session is None:
        raise ValueError("session is required for export_analysis_report_docx")
    record, markdown = _load_saved_report(analysis_run_id, session=session, data_dir=config_data_dir)
    root = _safe_output_dir(output_dir, data_dir=config_data_dir)
    output_path = _safe_output_path(root, _build_export_file_name(record.registry_number, record.id, "docx"))
    _build_docx(record, markdown, output_path)
    return _build_exported_report(record, output_path, "docx")


def export_analysis_report_pdf(
    analysis_run_id: str,
    *,
    output_dir: Path | None = None,
    data_dir: str | None = None,
    session: Session | None = None,
) -> ExportedReport:
    config_data_dir = data_dir or load_config().data_dir
    if session is None:
        raise ValueError("session is required for export_analysis_report_pdf")
    record, markdown = _load_saved_report(analysis_run_id, session=session, data_dir=config_data_dir)
    root = _safe_output_dir(output_dir, data_dir=config_data_dir)
    output_path = _safe_output_path(root, _build_export_file_name(record.registry_number, record.id, "pdf"))
    _build_pdf(record, markdown, output_path)
    return _build_exported_report(record, output_path, "pdf")


def get_exported_report_path(
    analysis_run_id: str,
    format: str,
    *,
    output_dir: Path | None = None,
    data_dir: str | None = None,
    session: Session | None = None,
) -> Path | None:
    format_name = format.lower()
    if format_name not in {"docx", "pdf"}:
        raise ValueError("Unsupported export format")
    config_data_dir = data_dir or load_config().data_dir
    if session is None:
        raise ValueError("session is required for get_exported_report_path")
    record = get_analysis_run(session, analysis_run_id)
    if record is None:
        return None
    root = _safe_output_dir(output_dir, data_dir=config_data_dir)
    candidate = _safe_output_path(root, _build_export_file_name(record.registry_number, record.id, format_name))
    return candidate if candidate.exists() else None


def list_report_exports(
    analysis_run_id: str,
    *,
    output_dir: Path | None = None,
    data_dir: str | None = None,
    session: Session | None = None,
) -> list[ExportedReport]:
    if session is None:
        raise ValueError("session is required for list_report_exports")
    record = get_analysis_run(session, analysis_run_id)
    if record is None:
        return []
    items: list[ExportedReport] = []
    for format_name in ("docx", "pdf"):
        path = get_exported_report_path(
            analysis_run_id,
            format_name,
            output_dir=output_dir,
            data_dir=data_dir,
            session=session,
        )
        if path is not None:
            items.append(_build_exported_report(record, path, format_name))
    return items
