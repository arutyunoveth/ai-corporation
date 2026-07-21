from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime, timezone
from html import escape as html_escape, unescape as html_unescape
from pathlib import Path
import hashlib
import json
import re
import os
import tempfile
from typing import Literal
from xml.sax.saxutils import escape as xml_escape

from docx import Document
from docx.shared import Pt
from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from src.modules.tender_operator_agent_demo.upload_service import (
    _load_metadata,
    get_uploaded_demo_report,
    get_uploaded_demo_report_html,
    get_demo_run_output_dir,
)
from src.tender_research.rag.export_service import (
    DOCX_CONTENT_TYPE,
    PDF_CONTENT_TYPE,
    _parse_markdown_blocks,
    _iter_inline_parts,
    _docx_add_formatted_paragraph,
    _ensure_pdf_font_registered,
    _pdf_styles,
    _pdf_inline_markup,
    _safe_segment,
    _short_run_id,
    _safe_output_path,
)


@dataclass(frozen=True)
class ExportedDemoReport:
    run_id: str
    registry_number: str
    format: Literal["docx", "pdf"]
    file_name: str
    file_path: str
    content_type: str
    size_bytes: int
    created_at: str
    source_report_path: str | None


_DEMO_EXPORT_SUBDIR = ("demo", "exports")
_SAFE_RUN_ID_RE = re.compile(r"^[A-Za-z0-9][A-Za-z0-9._-]{2,127}$")
_PDF_RENDERER_VERSION = "r7-persisted-pdf-v2"


def _analysis_title(registry_number: str) -> str:
    if registry_number:
        return f"Анализ закупки {registry_number}"
    return "Отчёт demo-agent"


def _validate_run_id(run_id: str) -> str:
    normalized = str(run_id or "").strip()
    if not normalized:
        raise ValueError("Run ID must not be empty")
    if "/" in normalized or "\\" in normalized or ".." in normalized:
        raise ValueError("Invalid run ID")
    if not _SAFE_RUN_ID_RE.fullmatch(normalized):
        raise ValueError("Invalid run ID")
    return normalized


def _demo_metadata_lines(metadata: dict) -> list[str]:
    procurement = metadata.get("procurement") or {}
    lines = [
        f"Реестровый номер: {metadata.get('procurement_id') or ''}",
        f"Закупка: {metadata.get('tender_title') or ''}",
        f"Заказчик: {metadata.get('customer_name') or ''}",
        f"Категория: {metadata.get('tender_category') or ''}",
        f"Статус: {metadata.get('status') or ''}",
        f"Режим анализа: {metadata.get('analysis_mode') or ''}",
        f"НМЦК: {_format_nmck(procurement.get('initial_price'), procurement.get('currency'))}",
        f"Дата публикации: {metadata.get('publication_date') or ''}",
        f"Срок подачи: {metadata.get('deadline') or ''}",
        f"Источник сведений: {procurement.get('structured_source_label') or metadata.get('notice_source_label') or 'карточка ЕИС'}",
    ]
    source = metadata.get("procurement_source") or ""
    if source:
        lines.append(f"Источник: {source}")
    created = metadata.get("created_at") or ""
    if created:
        if isinstance(created, datetime):
            created = created.isoformat()
        lines.append(f"Создано: {created}")
    return lines + [""]


def _format_nmck(price, currency) -> str:
    if price is None:
        return "не указана"
    try:
        val = float(price)
        formatted = f"{val:,.2f}".replace(",", " ")
    except (ValueError, TypeError):
        formatted = str(price)
    cur = currency or "RUB"
    return f"{formatted} {cur}"


def _build_export_file_name(registry_number: str, run_id: str, format_name: Literal["docx", "pdf"]) -> str:
    safe_registry = _safe_segment(registry_number, "unknown_registry")
    return f"demo_agent_report_{safe_registry}_{_short_run_id(run_id)}.{format_name}"


def _pdf_artifact_key(registry_number: str, run_id: str, report_model_hash: str) -> str:
    identity = f"{registry_number}\0{run_id}\0{report_model_hash}\0{_PDF_RENDERER_VERSION}".encode("utf-8")
    return hashlib.sha256(identity).hexdigest()[:24]


def _pdf_artifact_paths(root: Path, registry_number: str, run_id: str, report_model_hash: str) -> tuple[str, Path, Path]:
    key = _pdf_artifact_key(registry_number, run_id, report_model_hash)
    pdf = _safe_output_path(root, f"demo_agent_report_{_safe_segment(registry_number, 'unknown_registry')}_{key}.pdf")
    return key, pdf, pdf.with_suffix(".manifest.json")


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _validate_persisted_pdf(pdf_path: Path, manifest_path: Path, *, run_id: str, registry_number: str, report_model_hash: str, artifact_key: str) -> None:
    try:
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    except (OSError, ValueError, TypeError) as exc:
        raise RuntimeError("Persisted PDF manifest is missing or invalid") from exc
    expected = {
        "run_id": run_id,
        "registry_number": registry_number,
        "report_model_hash": report_model_hash,
        "artifact_key": artifact_key,
        "renderer_version": _PDF_RENDERER_VERSION,
    }
    if any(manifest.get(key) != value for key, value in expected.items()):
        raise RuntimeError("Persisted PDF manifest does not belong to requested run")
    if not pdf_path.is_file() or pdf_path.stat().st_size != manifest.get("byte_size"):
        raise RuntimeError("Persisted PDF artifact size is invalid")
    with pdf_path.open("rb") as existing:
        if existing.read(5) != b"%PDF-":
            raise RuntimeError("Persisted PDF artifact is corrupt")
    if _sha256(pdf_path) != manifest.get("pdf_sha256"):
        raise RuntimeError("Persisted PDF artifact hash is invalid")


def _write_pdf_manifest(path: Path, *, run_id: str, registry_number: str, report_model_hash: str, artifact_key: str, pdf_path: Path) -> None:
    payload = {"run_id": run_id, "registry_number": registry_number, "report_model_hash": report_model_hash, "artifact_key": artifact_key, "pdf_relative_path": f"data/demo/exports/{pdf_path.name}", "pdf_sha256": _sha256(pdf_path), "byte_size": pdf_path.stat().st_size, "renderer_version": _PDF_RENDERER_VERSION}
    temporary = path.with_suffix(path.suffix + ".partial")
    with temporary.open("w", encoding="utf-8") as handle:
        handle.write(json.dumps(payload, ensure_ascii=False, sort_keys=True) + "\n")
        handle.flush()
        os.fsync(handle.fileno())
    os.replace(temporary, path)


def _remove_interrupted_pdf_artifacts(pdf_path: Path, manifest_path: Path) -> None:
    """Discard only this key's incomplete/orphaned publication pair under its lock."""
    for path in (pdf_path, manifest_path, pdf_path.with_suffix(pdf_path.suffix + ".partial"), manifest_path.with_suffix(manifest_path.suffix + ".partial")):
        path.unlink(missing_ok=True)


def _fsync_directory(path: Path) -> None:
    try:
        descriptor = os.open(path, os.O_RDONLY)
        try:
            os.fsync(descriptor)
        finally:
            os.close(descriptor)
    except OSError:  # Filesystems such as some Docker mounts do not support directory fsync.
        pass


def _safe_output_dir(data_dir: str | None = None) -> Path:
    if data_dir:
        root = Path(data_dir).joinpath(*_DEMO_EXPORT_SUBDIR).resolve()
    else:
        root = Path.cwd().joinpath("data", *_DEMO_EXPORT_SUBDIR).resolve()
    root.mkdir(parents=True, exist_ok=True)
    return root


def _html_report_to_markdown(html: str) -> str:
    cleaned = re.sub(r"(?is)<(script|style)\b.*?>.*?</\1>", " ", html or "")
    replacements = [
        (r"(?i)<br\s*/?>", "\n"),
        (r"(?i)</p\s*>", "\n\n"),
        (r"(?i)</div\s*>", "\n"),
        (r"(?i)</li\s*>", "\n"),
        (r"(?i)<li\b[^>]*>", "- "),
        (r"(?i)</h1\s*>", "\n\n"),
        (r"(?i)</h2\s*>", "\n\n"),
        (r"(?i)</h3\s*>", "\n\n"),
        (r"(?i)<h1\b[^>]*>", "# "),
        (r"(?i)<h2\b[^>]*>", "## "),
        (r"(?i)<h3\b[^>]*>", "### "),
        (r"(?i)</tr\s*>", "\n"),
        (r"(?i)</td\s*>", " | "),
        (r"(?i)</th\s*>", " | "),
    ]
    for pattern, replacement in replacements:
        cleaned = re.sub(pattern, replacement, cleaned)
    cleaned = re.sub(r"(?is)<[^>]+>", " ", cleaned)
    cleaned = html_unescape(cleaned)
    cleaned = re.sub(r"[ \t]+\n", "\n", cleaned)
    cleaned = re.sub(r"\n[ \t]+", "\n", cleaned)
    cleaned = re.sub(r"[ \t]{2,}", " ", cleaned)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _report_markdown_for_export(run_id: str, report_markdown: str | None) -> str:
    if (report_markdown or "").strip():
        return str(report_markdown).strip()
    return _html_report_to_markdown(get_uploaded_demo_report_html(run_id))


def _build_docx_from_parts(
    title: str,
    metadata_lines: list[str],
    markdown: str,
    output_path: Path,
) -> None:
    document = Document()
    document.core_properties.title = title
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10.5)

    title_para = document.add_paragraph()
    title_para.style = document.styles["Title"]
    title_run = title_para.add_run(title)
    title_run.font.name = "Arial"

    for meta_line in metadata_lines:
        if meta_line:
            _docx_add_formatted_paragraph(document, meta_line)
        else:
            document.add_paragraph()

    document.add_heading("Отчёт", level=1)

    for block in _parse_markdown_blocks(markdown):
        if block.kind == "heading":
            document.add_heading(block.text, level=block.level)
        elif block.kind == "paragraph":
            _docx_add_formatted_paragraph(document, block.text)
        elif block.kind == "bullet_list":
            for item in block.items:
                _docx_add_formatted_paragraph(document, item, style="List Bullet")
        elif block.kind == "numbered_list":
            for item in block.items:
                _docx_add_formatted_paragraph(document, item, style="List Number")

    document.save(output_path)


def _build_pdf_from_parts(
    title: str,
    metadata_lines: list[str],
    markdown: str,
    output_path: Path,
) -> None:
    font_name = _ensure_pdf_font_registered()
    styles = _pdf_styles(font_name)
    story = [
        Paragraph(html_escape(title), styles["title"]),
        Spacer(1, 2 * mm),
    ]

    for meta_line in metadata_lines:
        if meta_line:
            story.append(Paragraph(_pdf_inline_markup(meta_line), styles["meta"]))
        else:
            story.append(Spacer(1, 2 * mm))

    story.append(Paragraph("Отчёт", styles["h1"]))

    for block in _parse_markdown_blocks(markdown):
        if block.kind == "heading":
            story.append(Paragraph(_pdf_inline_markup(block.text), styles[f"h{block.level}"]))
        elif block.kind == "paragraph":
            story.append(Paragraph(_pdf_inline_markup(block.text), styles["body"]))
        elif block.kind == "bullet_list":
            items = [ListItem(Paragraph(_pdf_inline_markup(item), styles["body"])) for item in block.items]
            story.append(ListFlowable(items, bulletType="bullet", start="circle"))
        elif block.kind == "numbered_list":
            items = [ListItem(Paragraph(_pdf_inline_markup(item), styles["body"])) for item in block.items]
            story.append(ListFlowable(items, bulletType="1"))

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
        title=title,
    )
    doc.build(story)


def _load_canonical_report(run_id: str) -> dict | None:
    path = get_demo_run_output_dir(run_id) / "canonical_report.json"
    if not path.is_file():
        return None
    import json
    return json.loads(path.read_text(encoding="utf-8"))


def _user_status(value: object) -> str:
    labels = {
        "known": "Объём определён",
        "partially_known": "Объём определён не для всех позиций",
        "specified": "Указано в документации",
        "absent": "Не обнаружено в документации",
        "extraction_failed": "Данные есть в источнике, но не извлечены автоматически",
        "not_specified": "Не обнаружено в документации",
        "not_specified_in_source": "Не обнаружено в документации",
        "not_applicable": "Товарный объём неприменим",
        "scope_conflict": "Тип закупки определён неоднозначно",
        "unknown": "Требуется проверка",
    }
    return labels.get(str(value), str(value))


def _user_evidence(row: dict) -> str:
    raw_document = str(row.get("source_document_id") or "Документ")
    lowered = raw_document.lower()
    if "xml" in lowered or "notification" in lowered:
        document = "Извещение ЕИС"
    elif "нмцк" in lowered:
        document = "Расчёт НМЦК"
    elif any(marker in lowered for marker in ("техническ", "ооз", "спецификац")):
        document = "Техническое задание"
    elif "контракт" in lowered:
        document = "Проект контракта"
    else:
        document = Path(raw_document).stem[:48] or "Документ"
    source_row = row.get("source_row")
    if source_row not in (None, ""):
        source_row = str(source_row).rsplit(":row:", 1)[-1].rsplit(":", 1)[-1]
    return f"{document}, позиция {source_row}" if source_row not in (None, "") else document


def _pdf_quantity(value: object) -> str:
    return "Не указано в проверенных документах" if value in {"Не указан документацией", "Не указано в проверенных документах"} else str(value)


def _build_docx_from_canonical(model: dict, title: str, output_path: Path) -> None:
    document = Document(); document.add_heading(title, 0)
    summary, passport = model["executive_summary"], model["procurement_passport"]
    contract_label = "приложен" if model.get("contract_draft_status") == "present" else ("приложен, но автоматически разобрать его не удалось" if model.get("contract_draft_status") == "parse_failed" else _user_status(model.get("contract_draft_status")))
    for line in (f"Официальное название закупки: {model.get('procurement_title')}", f"Номер закупки: {model.get('procurement_number')}", f"Заказчик: {model.get('customer_name')}", f"Место поставки: {model.get('delivery_place')}", f"Проект контракта: {contract_label}", f"Дата публикации: {model.get('publication_datetime')}", f"Окончание подачи заявок: {model.get('application_deadline')}", f"НМЦК: {model.get('nmck')} {model.get('currency')}", f"Решение: {model.get('decision')}", f"ОКПД2: {passport.get('okpd2')}", f"Статус объёма: {_user_status(model.get('procurement_volume_status'))}", f"Причина статуса объёма: {model.get('volume_status_reason')}"):
        document.add_paragraph(line)
    if model["line_items"]:
        document.add_paragraph(f"Первая извлечённая позиция: {model['line_items'][0].get('display_name') or model['line_items'][0]['original_name']}")
    document.add_heading("Состав и объём закупки", 1)
    table = document.add_table(rows=1, cols=6); table.style = "Table Grid"
    for cell, label in zip(table.rows[0].cells, ("№", "Наименование", "Количество", "Единица", "Статус количества", "Источник / evidence")): cell.text = label
    for row in model["line_items"]:
        cells = table.add_row().cells
        values = (row["sequence"], row.get("display_name") or row["original_name"], row["quantity_display"], row["unit_original"], _user_status(row["quantity_status"]), _user_evidence(row))
        for cell, value in zip(cells, values): cell.text = str(value)
    if not model["line_items"]:
        document.add_paragraph("Позиции и количество не удалось извлечь из доступных документов; требуется проверка первоисточника.")
    document.add_heading("Ограничения анализа", 1)
    for item in model["missing_data"]: document.add_paragraph(item["description"], style="List Bullet")
    for item in model["limitations"]: document.add_paragraph(item, style="List Bullet")
    document.add_paragraph("Техническая карта доказательств сохранена в JSON-версии отчёта.")
    document.save(output_path)


def _build_pdf_from_canonical(model: dict, title: str, output_path: Path) -> None:
    font_name = _ensure_pdf_font_registered(); styles = _pdf_styles(font_name)
    table_style = ParagraphStyle("focused_table", parent=styles["body"], fontSize=6.7, leading=7.5)
    summary, passport = model["executive_summary"], model["procurement_passport"]
    story = [Paragraph(html_escape(title), styles["title"])]
    contract_label = "приложен" if model.get("contract_draft_status") == "present" else ("приложен, но автоматически разобрать его не удалось" if model.get("contract_draft_status") == "parse_failed" else _user_status(model.get("contract_draft_status")))
    story += [Paragraph(_pdf_inline_markup(str(line)), styles["body"]) for line in (f"Название закупки: {model.get('procurement_title')}", f"Номер закупки: {model.get('procurement_number')}", f"Заказчик: {model.get('customer_name')}", f"Место поставки: {model.get('delivery_place')}", f"Проект контракта: {contract_label}", f"Дата публикации: {model.get('publication_datetime')}", f"Окончание подачи заявок: {model.get('application_deadline')}", f"НМЦК: {model.get('nmck')} {model.get('currency')}", f"Решение: {model.get('decision')}", f"ОКПД2: {passport.get('okpd2')}", f"Статус объёма: {_user_status(model.get('procurement_volume_status'))}")]
    story.append(Paragraph("Состав и объём закупки", styles["h1"]))
    source_issues = [issue for row in model.get("line_items", []) for issue in row.get("field_issues", [])]
    if source_issues:
        story.append(Paragraph("Подтверждённые ограничения исходных документов", styles["h1"]))
        okpd2_missing = sum(issue.get("field_name") == "okpd2" for issue in source_issues)
        quantity_missing = sum(issue.get("field_name") == "quantity" for issue in source_issues)
        total_items = len(model.get("line_items", []))
        if okpd2_missing:
            story.append(Paragraph(f"Для {okpd2_missing} из {total_items} позиций код ОКПД2 не найден в проверенных документах. Позиции идентифицированы по наименованию и месту в источнике.", styles["body"]))
        if quantity_missing:
            story.append(Paragraph(f"Количество не найдено для {quantity_missing} позиций. Коммерческий расчёт требует ручной проверки.", styles["body"]))
    # Put limitations before the variable-height table.  This prevents a
    # short limitations-only tail page after a table that happened to fit.
    story.append(Paragraph("Ограничения анализа", styles["h1"]))
    story += [Paragraph(_pdf_inline_markup(item["description"]), styles["body"]) for item in model["missing_data"]]
    story += [Paragraph(_pdf_inline_markup(item), styles["body"]) for item in model["limitations"]]
    data = [[Paragraph(value, table_style) for value in ("№", "Наименование", "Количество", "Ед.", "Статус", "Источник / evidence")]]
    for row in model["line_items"]:
        characteristics = row.get("characteristics") or []
        characteristic_text = "; ".join(str(item.get("display_value") or "") for item in characteristics[:3] if isinstance(item, dict))
        display_name = row.get("display_name") or row["original_name"]
        item_cell = [Paragraph(_pdf_inline_markup(str(display_name)), table_style)]
        if characteristic_text:
            item_cell.append(Paragraph(_pdf_inline_markup(characteristic_text), ParagraphStyle("focused_characteristics", parent=table_style, fontSize=6, leading=7)))
        row_values = (row["sequence"], item_cell, _pdf_quantity(row["quantity_display"]), row["unit_original"], _user_status(row["quantity_status"]), _user_evidence(row))
        data.append([value if isinstance(value, list) else Paragraph(_pdf_inline_markup(str(value)), table_style) for value in row_values])
    if len(data) == 1:
        empty_message = "Основной предмет закупки — работы. Товарный анализ неприменим." if model.get("procurement_scope", {}).get("procurement_primary_scope") == "works" else "Позиции и количество не удалось извлечь из доступных документов; требуется проверка первоисточника."
        story.append(Paragraph(empty_message, styles["body"]))
    else:
        table = Table(data, colWidths=[8*mm, 58*mm, 25*mm, 22*mm, 30*mm, 27*mm], repeatRows=1)
        table.setStyle(TableStyle([("GRID", (0,0), (-1,-1), .25, HexColor("#b9c8d0")), ("BACKGROUND", (0,0), (-1,0), HexColor("#e9f7f5")), ("VALIGN", (0,0), (-1,-1), "TOP")]))
        story.append(table)
    model_hash = ((model.get("provenance") or {}).get("production_model_hash") or model.get("production_model_hash") or "unknown")
    doc = SimpleDocTemplate(str(output_path), pagesize=A4, leftMargin=12*mm, rightMargin=12*mm, topMargin=14*mm, bottomMargin=14*mm, title=title, author=f"production_model_hash={model_hash}"); doc.build(story)


def export_demo_agent_report_docx(run_id: str) -> ExportedDemoReport:
    run_id = _validate_run_id(run_id)
    metadata = _load_metadata(run_id)
    report = get_uploaded_demo_report(run_id)
    report_markdown = _report_markdown_for_export(run_id, report.report_markdown)

    registry_number = metadata.get("procurement_id") or metadata.get("reestr_number") or ""
    title = _analysis_title(registry_number)
    metadata_lines = _demo_metadata_lines(metadata)

    root = _safe_output_dir()
    file_name = _build_export_file_name(registry_number, run_id, "docx")
    output_path = _safe_output_path(root, file_name)
    canonical = _load_canonical_report(run_id)
    if canonical:
        _build_docx_from_canonical(canonical, title, output_path)
    else:
        _build_docx_from_parts(title, metadata_lines, report_markdown, output_path)

    return ExportedDemoReport(
        run_id=run_id,
        registry_number=registry_number,
        format="docx",
        file_name=file_name,
        file_path=str(output_path),
        content_type=DOCX_CONTENT_TYPE,
        size_bytes=output_path.stat().st_size,
        created_at=datetime.now(timezone.utc).isoformat(),
        source_report_path=str(metadata.get("_metadata_path", "")),
    )


def export_demo_agent_report_pdf(run_id: str) -> ExportedDemoReport:
    run_id = _validate_run_id(run_id)
    metadata = _load_metadata(run_id)
    report = get_uploaded_demo_report(run_id)
    report_markdown = _report_markdown_for_export(run_id, report.report_markdown)

    registry_number = metadata.get("procurement_id") or metadata.get("reestr_number") or ""
    title = _analysis_title(registry_number)
    metadata_lines = _demo_metadata_lines(metadata)

    canonical = _load_canonical_report(run_id)
    report_model_hash = hashlib.sha256(json.dumps(canonical, ensure_ascii=False, sort_keys=True).encode("utf-8")).hexdigest() if canonical else hashlib.sha256(report_markdown.encode("utf-8")).hexdigest()
    root = _safe_output_dir()
    artifact_key, output_path, manifest_path = _pdf_artifact_paths(root, registry_number, run_id, report_model_hash)
    file_name = output_path.name
    # A completed run's customer PDF is an immutable persisted artifact.
    lock_path = output_path.with_suffix(output_path.suffix + ".lock")
    with lock_path.open("a+b") as lock:
        try:
            import fcntl
            fcntl.flock(lock.fileno(), fcntl.LOCK_EX)
        except ImportError:  # pragma: no cover - Windows development fallback
            pass
        # Partial files never represent a published artifact.  A lone member of
        # the pair is recoverable after a process crash, not a permanent error.
        for partial in (output_path.with_suffix(output_path.suffix + ".partial"), manifest_path.with_suffix(manifest_path.suffix + ".partial")):
            partial.unlink(missing_ok=True)
        if output_path.is_file() and manifest_path.is_file():
            _validate_persisted_pdf(output_path, manifest_path, run_id=run_id, registry_number=registry_number, report_model_hash=report_model_hash, artifact_key=artifact_key)
        else:
            if output_path.is_file() or manifest_path.is_file():
                _remove_interrupted_pdf_artifacts(output_path, manifest_path)
            descriptor, temporary_name = tempfile.mkstemp(prefix=".pdf-", suffix=".partial", dir=root)
            os.close(descriptor)
            temporary_path = Path(temporary_name)
            try:
                if canonical:
                    _build_pdf_from_canonical(canonical, title, temporary_path)
                else:
                    _build_pdf_from_parts(title, metadata_lines, report_markdown, temporary_path)
                with temporary_path.open("rb") as created:
                    if created.read(5) != b"%PDF-":
                        raise RuntimeError("Generated PDF artifact is invalid")
                    created.seek(0, os.SEEK_END)
                    if created.tell() == 0:
                        raise RuntimeError("Generated PDF artifact is empty")
                with temporary_path.open("rb") as created:
                    os.fsync(created.fileno())
                os.replace(temporary_path, output_path)
                _write_pdf_manifest(manifest_path, run_id=run_id, registry_number=registry_number, report_model_hash=report_model_hash, artifact_key=artifact_key, pdf_path=output_path)
                _fsync_directory(root)
            finally:
                temporary_path.unlink(missing_ok=True)

    return ExportedDemoReport(
        run_id=run_id,
        registry_number=registry_number,
        format="pdf",
        file_name=file_name,
        file_path=str(output_path),
        content_type=PDF_CONTENT_TYPE,
        size_bytes=output_path.stat().st_size,
        created_at=datetime.now(timezone.utc).isoformat(),
        source_report_path=str(metadata.get("_metadata_path", "")),
    )
